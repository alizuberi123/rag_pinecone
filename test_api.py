import os
import io
import re
from typing import List, Optional, Dict, Any
from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import pdfplumber
import docx
import pytesseract
from PIL import Image
import fitz  # PyMuPDF
import google.generativeai as genai
from google.generativeai.types import content_types
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from dotenv import load_dotenv
from datetime import datetime

from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from utils.rag_manager import create_rag, list_rags, get_rag_path, save_chat_history, load_chat_history
from utils.retrieval import HybridRetriever

# Load environment variables
load_dotenv()
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# Use DATA_DIR env variable for base data directory
BASE_DATA_DIR = os.environ.get("DATA_DIR", ".")
RAG_SESSIONS_DIR = os.path.join(BASE_DATA_DIR, "rag_sessions")
active_rag = "default"  # You can modify this based on your needs
rag_path = os.path.join(RAG_SESSIONS_DIR, active_rag)
db_path = os.path.join(rag_path, "chroma_db")
file_dir = os.path.join(rag_path, "files")

app = FastAPI()

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Pydantic models for request/response
class QueryRequest(BaseModel):
    query: str
    car_make: str
    car_model: str
    car_year: str

class QueryResponse(BaseModel):
    response: str
    topic: str
    sources: List[Dict[str, Any]]
    diagrams: List[str] = []

class UploadResponse(BaseModel):
    message: str
    filename: str
    text_chunks: int
    diagram_chunks: int
    car_make: str
    car_model: str
    car_year: str

class MemorySettings(BaseModel):
    max_tokens: int = 2000
    max_messages: int = 10
    include_summary: bool = True

class ThreadRequest(BaseModel):
    thread_id: str
    settings: Optional[MemorySettings] = None

class ConversationMemory:
    def __init__(self, max_tokens: int = 2000, max_messages: int = 10, include_summary: bool = True):
        self.max_tokens = max_tokens
        self.max_messages = max_messages
        self.include_summary = include_summary
        self.threads: Dict[str, List[Dict[str, Any]]] = {}
        self.summaries: Dict[str, str] = {}

    def add_message(self, thread_id: str, query: str, response: str, sources: List[Dict[str, Any]]):
        if thread_id not in self.threads:
            self.threads[thread_id] = []
            self.summaries[thread_id] = ""

        message = {
            "query": query,
            "response": response,
            "sources": sources,
            "timestamp": datetime.now().isoformat()
        }

        self.threads[thread_id].append(message)

        # Trim old messages if needed
        if len(self.threads[thread_id]) > self.max_messages:
            self.threads[thread_id] = self.threads[thread_id][-self.max_messages:]

        # Update summary if enabled
        if self.include_summary:
            self._update_summary(thread_id)

    def _update_summary(self, thread_id: str):
        if not self.threads[thread_id]:
            return

        model = genai.GenerativeModel('models/gemini-1.5-pro-latest')
        conversation = "\n\n".join([
            f"Q: {msg['query']}\nA: {msg['response']}"
            for msg in self.threads[thread_id]
        ])

        try:
            summary = model.generate_content(
                f"""Summarize the key points from this conversation in a concise way:

                {conversation}

                Focus on the main topics discussed and any important conclusions or decisions made.
                Keep the summary under 200 words."""
            ).text
            self.summaries[thread_id] = summary
        except:
            pass

    def get_context(self, thread_id: str) -> str:
        if thread_id not in self.threads:
            return ""

        context_parts = []
        if self.include_summary and self.summaries[thread_id]:
            context_parts.append(f"Conversation Summary:\n{self.summaries[thread_id]}\n")

        recent_messages = self.threads[thread_id][-3:]  # Get last 3 messages
        for msg in recent_messages:
            context_parts.append(f"Previous Q: {msg['query']}\nPrevious A: {msg['response']}\n")

        return "\n".join(context_parts)

# Initialize global variables
embed_model = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
active_rag = "default"  # You can modify this based on your needs
rag_path = get_rag_path(active_rag)
db_path = os.path.join(rag_path, "chroma_db")
file_dir = os.path.join(rag_path, "files")

# Initialize Chroma DB
db = Chroma(
    persist_directory=db_path,
    embedding_function=embed_model,
    collection_name="document_collection"
)

# Initialize retriever and memory
retriever = HybridRetriever(db, embed_model)
memory = ConversationMemory()

def clean_text(text):
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove special characters
    text = re.sub(r'[^\w\s.,!?-]', '', text)
    return text.strip()

def extract_text_and_diagrams(file_path: str, file_type: str):
    text_chunks = []
    diagram_chunks = []

    if file_type == "application/pdf":
        with pdfplumber.open(file_path) as pdf:
            # Extract section headings for each page
            page_sections = []
            section_pattern = re.compile(r'^(\d+(\.\d+)*\.?\s+.+|[A-Z][A-Z\s]{3,})$', re.MULTILINE)
            for page in pdf.pages:
                text = page.extract_text() or ""
                # Try to find a section heading in the first 10 lines
                lines = text.split('\n')[:10]
                section = "General"
                for line in lines:
                    if section_pattern.match(line.strip()):
                        section = line.strip()
                        break
                page_sections.append(section)
            # Now extract text and assign section
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                cleaned = clean_text(text)
                if cleaned:
                    text_chunks.append({
                        "text": cleaned,
                        "section": page_sections[i],
                        "page_number": i + 1
                    })
        # Extract diagrams and assign section
        pdf_doc = fitz.open(file_path)
        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            image_list = page.get_images(full=True)
            section = "General"
            if page_num < len(page_sections):
                section = page_sections[page_num]
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = pdf_doc.extract_image(xref)
                image_bytes = base_image["image"]
                img_pil = Image.open(io.BytesIO(image_bytes)).convert("RGB")
                image_dir = os.path.join(file_dir, "diagrams")
                os.makedirs(image_dir, exist_ok=True)
                image_path = os.path.join(image_dir, f"{os.path.basename(file_path)}_page{page_num}_{img_index}.png")
                img_pil.save(image_path)
                ocr_text = pytesseract.image_to_string(img_pil)
                cleaned = clean_text(ocr_text)
                if not cleaned.strip():
                    try:
                        vision_model = genai.GenerativeModel("models/gemini-1.5-pro-latest")
                        image_data = content_types.ImageData.from_pil_image(img_pil)
                        gemini_result = vision_model.generate_content(
                            [image_data, "Describe the diagram in detail as if explaining it to a reader."]
                        )
                        cleaned = gemini_result.text.strip()
                    except Exception as e:
                        print(f"Error processing diagram: {str(e)}")
                        cleaned = "Unrecognized diagram content"
                if cleaned:
                    diagram_chunks.append({
                        "text": cleaned,
                        "source": os.path.basename(file_path),
                        "page": page_num + 1,
                        "index": img_index,
                        "image_path": image_path,
                        "section": section
                    })
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        text_chunks.append({"text": clean_text(text), "section": "General", "page_number": 1})
    elif file_type == "text/plain":
        with open(file_path, 'r') as f:
            text = f.read()
            text_chunks.append({"text": clean_text(text), "section": "General", "page_number": 1})
    return text_chunks, diagram_chunks

def structured_chunking(raw_text: str, filename: str):
    base_name = filename.lower().replace(".pdf", "").replace(".docx", "").replace(".txt", "")
    name_parts = re.split(r"[_\-\s]+", base_name)
    tags = [part.capitalize() for part in name_parts if part]

    section_pattern = re.compile(r'^(\d+(\.\d+)*\.?\s+.+|[A-Z][A-Z\s]{3,})$', re.MULTILINE)
    lines = raw_text.split('\n')
    sections = []
    current_section = "General"
    current_content = []

    for line in lines:
        if section_pattern.match(line.strip()):
            if current_content:
                sections.append((current_section, '\n'.join(current_content)))
                current_content = []
            current_section = line.strip()
        else:
            current_content.append(line.strip())

    if current_content:
        sections.append((current_section, '\n'.join(current_content)))

    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    documents = []

    for section_title, section_text in sections:
        chunks = splitter.split_text(section_text)
        for i, chunk in enumerate(chunks):
            documents.append(Document(
                page_content=chunk.strip(),
                metadata={
                    "filename": filename,
                    "section": section_title.strip(),
                    "chunk_id": i,
                    "tags": ", ".join(tags),
                    "type": "text"
                }
            ))
    return documents

def infer_topic_from_prompt(prompt: str) -> str:
    model = genai.GenerativeModel('models/gemini-1.5-pro-latest')
    try:
        upgraded_prompt = f"""
You are a domain expert categorizer.

Given the following user question, identify the most specific technical topic or document section it relates to, based only on the wording of the question.

Question:
{prompt}

Respond with just the topic name. Do not include any explanations or extra text.
"""
        return model.generate_content(upgraded_prompt).text.strip()
    except:
        return "General"

def compress_chunks(chunks: List[Document], question: str) -> str:
    model = genai.GenerativeModel('models/gemini-1.5-pro-latest')
    combined = "\n\n".join([doc.page_content.strip() for doc in chunks])
    
    prompt = f"""
You are summarizing content for a technical Q&A assistant that answers user questions using uploaded documents.

Here is the user's question:
{question}

And here are the top most relevant document chunks:
{combined}

Your task:
- Combine the content into a clean, concise summary focused on answering the user's question.
- Remove duplicate sentences, repeated facts, or off-topic remarks.
- Keep essential details, technical terms, and step-by-step explanations.
- If applicable, include short quotes or bullet points.
- Keep the summary under 300 words if possible.

Respond with the cleaned and compressed version of the information only. Do not restate the question or explain what you're doing.
"""
    try:
        return model.generate_content(prompt).text.strip()
    except:
        return combined

def extract_car_info(filename: str, car_make: str = None, car_model: str = None, car_year: str = None):
    # Try to extract from filename if not provided
    if not (car_make and car_model and car_year):
        import re
        # Example: Toyota_Camry_2018.pdf
        match = re.match(r"([A-Za-z]+)_([A-Za-z0-9]+)_([0-9]{4})", filename)
        if match:
            car_make = car_make or match.group(1)
            car_model = car_model or match.group(2)
            car_year = car_year or match.group(3)
    return car_make, car_model, car_year

@app.get("/")
async def root():
    return {"message": "Hello World"}

@app.get("/health")
async def health_check():
    return {"status": "healthy"}

@app.get("/rags")
async def list_rag_sessions():
    return {"rags": list_rags()}

@app.post("/create_rag/{name}")
async def create_rag_session(name: str):
    if create_rag(name):
        return {"message": f"RAG session '{name}' created successfully"}
    raise HTTPException(status_code=400, detail="RAG session already exists")

@app.post("/upload", response_model=UploadResponse)
async def upload_file(
    file: UploadFile = File(...),
    car_make: str = Form(None),
    car_model: str = Form(None),
    car_year: str = Form(None)
):
    try:
        file_path = os.path.join(file_dir, file.filename)
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        car_make, car_model, car_year = extract_car_info(file.filename, car_make, car_model, car_year)
        if not (car_make and car_model and car_year):
            raise HTTPException(status_code=400, detail="car_make, car_model, and car_year are required (either as form fields or in the filename)")
        text_chunks, diagram_chunks = extract_text_and_diagrams(file_path, file.content_type)
        documents = []
        for chunk in text_chunks:
            docs = structured_chunking(chunk["text"], file.filename)
            for doc in docs:
                doc.metadata["car_make"] = car_make
                doc.metadata["car_model"] = car_model
                doc.metadata["car_year"] = car_year
                doc.metadata["section"] = chunk["section"]
                doc.metadata["page_number"] = chunk["page_number"]
            documents.extend(docs)
        if documents:
            db.add_documents(documents)
        for diagram in diagram_chunks:
            doc = Document(
                page_content=diagram["text"],
                metadata={
                    "filename": diagram["source"],
                    "section": diagram["section"],
                    "type": "diagram",
                    "image_path": diagram["image_path"],
                    "car_make": car_make,
                    "car_model": car_model,
                    "car_year": car_year,
                    "page_number": diagram["page"]
                }
            )
            db.add_documents([doc])
        return UploadResponse(
            message="File processed successfully",
            filename=file.filename,
            text_chunks=len(documents),
            diagram_chunks=len(diagram_chunks),
            car_make=car_make,
            car_model=car_model,
            car_year=car_year
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/query", response_model=QueryResponse)
async def query_documents(query: QueryRequest):
    try:
        topic = infer_topic_from_prompt(query.query)
        chunks = retriever.get_relevant_documents(query.query, k=15)
        # Filter by car metadata
        filtered_chunks = [doc for doc in chunks if (
            doc.metadata.get("car_make", "").lower() == query.car_make.lower() and
            doc.metadata.get("car_model", "").lower() == query.car_model.lower() and
            str(doc.metadata.get("car_year", "")) == str(query.car_year)
        )]
        if not filtered_chunks:
            raise HTTPException(status_code=404, detail="No relevant information found for the specified car.")
        filtered_chunks = filtered_chunks[:5]
        # Get the set of relevant sections from the text chunks
        relevant_sections = set(doc.metadata.get("section") for doc in filtered_chunks if doc.metadata.get("type", "text") == "text")
        # Only include diagrams whose section matches a relevant section
        diagram_paths = [doc.metadata.get("image_path") for doc in filtered_chunks if doc.metadata.get("type") == "diagram" and doc.metadata.get("image_path") and doc.metadata.get("section") in relevant_sections]
        compressed_content = compress_chunks(filtered_chunks, query.query)
        model = genai.GenerativeModel('models/gemini-1.5-pro-latest')
        prompt = f"""
You are a highly knowledgeable automotive assistant. Answer the user's question using the provided car manual context below. 
If the manual is vague or incomplete, supplement the answer with your own expert knowledge, but always be specific, step-by-step, and accurate for the given car make, model, and year.

Context from the manual:
{compressed_content}

Question: {query.query}

Answer:
"""
        response = model.generate_content(prompt)
        return QueryResponse(
            response=response.text,
            topic=topic,
            sources=[doc.metadata for doc in filtered_chunks],
            diagrams=diagram_paths
        )
    except HTTPException as e:
        raise e
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Internal error: {str(e)}") 